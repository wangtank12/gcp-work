from openpyxl import load_workbook
from docx import Document
from datetime import datetime
from io import BytesIO
import os
import shutil
from docxcompose.composer import Composer

def replace_single_row_data():
    # 定义输入路径
    excel_path = r"D:\工作\问询表\项目数据表.xlsx"
    template_path = r"D:\工作\问询表\受试者问询记录表.docx"
    
    # 加载工作簿和活动工作表
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    
    # 获取数据名称和数据内容
    headers = [cell.value for cell in ws[1]]
    data = {}
    
    # 处理单行数据
    for i, header in enumerate(headers):
        cell_value = ws.cell(row=2, column=i+1).value
        if cell_value and not ws.cell(row=3, column=i+1).value:
            # 如果是日期类型，去除时间部分
            if isinstance(cell_value, datetime):
                data[header] = cell_value.date()  # 仅保留日期部分
            else:
                data[header] = cell_value
    
    # 创建动态输出文件名
    project_name = data.get('项目名称', '')
    project_stage = data.get('阶段', '')
    original_name = os.path.splitext(os.path.basename(template_path))[0]
    output_name = f"{original_name}_{project_name}{project_stage}.docx"
    output_path = os.path.join(r"D:\工作\问询表", output_name)
    
    # 替换Word模板中的占位符
    doc = Document(template_path)
    
    # 处理普通段落（修复字体复制问题）
    for paragraph in doc.paragraphs:
        for header, value in data.items():
            placeholder = f"--{header}--"
            if placeholder in paragraph.text:
                original_runs = [(run.text, run.font) for run in paragraph.runs]
                paragraph.clear()
                
                for text, font in original_runs:
                    new_run = paragraph.add_run(text.replace(placeholder, str(value)))
                    # 手动复制关键字体属性
                    new_run.font.name = font.name
                    new_run.font.size = font.size
                    new_run.font.bold = font.bold
                    new_run.font.italic = font.italic

    # 处理表格中的占位符（同样修复）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for header, value in data.items():
                        placeholder = f"--{header}--"
                        if placeholder in paragraph.text:
                            original_runs = [(run.text, run.font) for run in paragraph.runs]
                            paragraph.clear()
                            
                            for text, font in original_runs:
                                new_run = paragraph.add_run(text.replace(placeholder, str(value)))
                                # 手动复制关键字体属性
                                new_run.font.name = font.name
                                new_run.font.size = font.size
                                new_run.font.bold = font.bold
                                new_run.font.italic = font.italic
    
    # 替换为兼容WPS的保存方式
    mem_stream = BytesIO()
    doc.save(mem_stream)
    mem_stream.seek(0)  # 确保流指针回到起始位置
    fixed_doc = Document(mem_stream)
    
    # 添加清理操作
    del doc
    mem_stream.close()
    
    fixed_doc.save(output_path)
    return output_path

def generate_multiple_docs(template_path):
    # 定义输入输出路径
    excel_path = r"D:\工作\问询表\项目数据表.xlsx"
    output_dir = r"D:\工作\问询表\输出文档"
    
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # 加载Excel数据
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    
    # 找出有多行的列
    multi_row_headers = []
    for i, header in enumerate(headers):
        if ws.cell(row=3, column=i+1).value:
            multi_row_headers.append(header)
    
    # 邮件合并生成多个文档
    for row_num in range(2, ws.max_row + 1):
        doc = Document(template_path)
        for header in multi_row_headers:
            col_index = headers.index(header) + 1
            cell_value = ws.cell(row=row_num, column=col_index).value
            if cell_value:
                # 处理日期格式
                if isinstance(cell_value, datetime):
                    value = cell_value.date()  # 仅保留日期部分
                else:
                    value = cell_value
                
                placeholder = f"--{header}--"
                
                # 处理普通段落
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, str(value))
                
                # 处理表格内容
                for table in doc.tables:
                    for table_row in table.rows:
                        for cell in table_row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(placeholder, str(value))
        
        # 保存生成的文档
        output_path = os.path.join(output_dir, f"受试者问询记录_{row_num-1}.docx")
        doc.save(output_path)
    
    # 修改后的合并文档函数
    def merge_documents():
        merged_path = os.path.join(output_dir, "合并文档.docx")
        # 获取所有生成的文档文件并按序号排序
        doc_files = sorted(
            [f for f in os.listdir(output_dir) if f.startswith("受试者问询记录_") and f.endswith(".docx")],
            key=lambda x: int(x.split('_')[1].split('.')[0])
        )
        
        if not doc_files:
            return None

        # 创建合并文档
        master = Document(os.path.join(output_dir, doc_files[0]))
        composer = Composer(master)
        for file in doc_files[1:]:
            doc_path = os.path.join(output_dir, file)
            doc = Document(doc_path)
            composer.append(doc)
        composer.save(merged_path)

        # 读取Excel数据构建最终文件名
        project_name = ws.cell(row=2, column=2).value  # 第二列
        project_stage = ws.cell(row=2, column=4).value  # 第四列
        template_base = os.path.splitext(os.path.basename(template_path))[0]
        
        # 移除可能重复的项目信息
        if template_base.endswith(project_name + project_stage):
            template_base = template_base.replace(project_name + project_stage, "")
        
        new_name = f"{template_base}{project_name}{project_stage}_合并版.docx"
        final_path = os.path.join(os.path.dirname(output_dir), new_name)

        # 移动文件到上级目录
        shutil.move(merged_path, final_path)
        return final_path

    # 执行合并并获取最终路径
    final_path = merge_documents()

    # 清理操作
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)

    return final_path

if __name__ == "__main__":
    # 第一步：创建新模板
    new_template_path = replace_single_row_data()
    print(f"已创建新模板: {new_template_path}")
    
    # 第二步：生成多个文档并合并
    final_output = generate_multiple_docs(new_template_path)
    
    # 第三步：清理新模板文件
    if os.path.exists(new_template_path):
        os.remove(new_template_path)
    
    print(f"文档生成完成！最终合并文档路径：{final_output}")