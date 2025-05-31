from openpyxl import load_workbook
from docx import Document  # 全局导入
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from datetime import datetime  # 新增导入
from io import BytesIO
import os
from docxcompose.composer import Composer  # 新增导入

def replace_single_row_data():
    # 将路径定义移动到函数开头
    excel_path = r"D:\工作\问询表\项目数据表.xlsx"
    template_path = r"D:\工作\问询表\受试者问询记录表.docx"
    output_path = r"D:\工作\问询表\新模板.docx"
    
    # 然后再进行其他操作
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    
    # 获取数据名称和数据内容
    headers = [cell.value for cell in ws[1]]
    data = {}
    
    # 处理单行数据
    for i, header in enumerate(headers):
        cell_value = ws.cell(row=2, column=i+1).value
        if cell_value and not ws.cell(row=3, column=i+1).value:
            # 如果是日期类型，去除时间部分
            if isinstance(cell_value, datetime):
                data[header] = cell_value.date()  # 仅保留日期部分
            else:
                data[header] = cell_value
    
    # 替换Word模板中的占位符
    doc = Document(template_path)
    
    # 处理普通段落（修复字体复制问题）
    for paragraph in doc.paragraphs:
        for header, value in data.items():
            placeholder = f"--{header}--"
            if placeholder in paragraph.text:
                original_runs = [(run.text, run.font) for run in paragraph.runs]
                paragraph.clear()
                
                for text, font in original_runs:
                    new_run = paragraph.add_run(text.replace(placeholder, str(value)))
                    # 手动复制关键字体属性
                    new_run.font.name = font.name
                    new_run.font.size = font.size
                    new_run.font.bold = font.bold
                    new_run.font.italic = font.italic

    # 处理表格中的占位符（同样修复）
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for header, value in data.items():
                        placeholder = f"--{header}--"
                        if placeholder in paragraph.text:
                            original_runs = [(run.text, run.font) for run in paragraph.runs]
                            paragraph.clear()
                            
                            for text, font in original_runs:
                                new_run = paragraph.add_run(text.replace(placeholder, str(value)))
                                # 手动复制关键字体属性
                                new_run.font.name = font.name
                                new_run.font.size = font.size
                                new_run.font.bold = font.bold
                                new_run.font.italic = font.italic
    
    # 替换为兼容WPS的保存方式
    mem_stream = BytesIO()
    doc.save(mem_stream)
    mem_stream.seek(0)  # 确保流指针回到起始位置
    fixed_doc = Document(mem_stream)
    
    # 添加清理操作
    del doc
    mem_stream.close()
    
    fixed_doc.save(output_path)
    return output_path

def generate_multiple_docs():
    # 添加路径定义（与replace_single_row_data中相同）
    excel_path = r"D:\工作\问询表\项目数据表.xlsx"
    template_path = r"D:\工作\问询表\新模板.docx"
    output_dir = r"D:\工作\问询表\输出文档"
    
    # 原有代码保持不变
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    wb = load_workbook(excel_path, data_only=True)
    ws = wb.active
    
    # 获取数据名称
    headers = [cell.value for cell in ws[1]]
    
    # 找出有多行的列
    multi_row_headers = []
    for i, header in enumerate(headers):
        if ws.cell(row=3, column=i+1).value:
            multi_row_headers.append(header)
    
    # 邮件合并生成多个文档
    for row_num in range(2, ws.max_row + 1):  # 修改循环变量名为row_num
        doc = Document(template_path)
        for header in multi_row_headers:
            col_index = headers.index(header) + 1
            cell_value = ws.cell(row=row_num, column=col_index).value
            if cell_value:
                # 处理日期格式
                if isinstance(cell_value, datetime):
                    value = cell_value.date()  # 仅保留日期部分
                else:
                    value = cell_value
                
                placeholder = f"--{header}--"
                
                # 处理普通段落
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(placeholder, str(value))
                
                # 处理表格内容
                for table in doc.tables:
                    for table_row in table.rows:  # 修改内部循环变量名
                        for cell in table_row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder in paragraph.text:
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(placeholder, str(value))
        
        # 保存生成的文档
        output_path = os.path.join(output_dir, f"受试者问询记录_{row_num-1}.docx")  # 使用row_num
        # 修改文档保存方式
        mem_stream = BytesIO()
        doc.save(mem_stream)
        mem_stream.seek(0)  # 确保流指针回到起始位置
        fixed_doc = Document(mem_stream)
        
        # 添加清理操作
        del doc
        mem_stream.close()
        
        fixed_doc.save(output_path)
    
    # 在函数末尾添加合并文档功能
    def merge_documents():
        merged_path = os.path.join(output_dir, "合并文档.docx")
        master = Document(os.path.join(output_dir, "受试者问询记录_1.docx"))
        composer = Composer(master)
        
        for i in range(2, ws.max_row):
            doc_path = os.path.join(output_dir, f"受试者问询记录_{i}.docx")
            if os.path.exists(doc_path):
                doc = Document(doc_path)
                composer.append(doc)
        
        composer.save(merged_path)
        return merged_path
    
    merge_documents()  # 执行合并
    return output_dir

if __name__ == "__main__":
    new_template = replace_single_row_data()
    output_folder = generate_multiple_docs()
    print(f"文档生成完成！单个文档保存在：{output_folder}，合并文档已生成")